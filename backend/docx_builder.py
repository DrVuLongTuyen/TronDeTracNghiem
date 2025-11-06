import docx
import json
import random 
import io       
import zipfile  
from docx import Document 
from docx.shared import Pt, Cm 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re 

# (MỚI V26) Import từ file utils
from docx_utils import style_run, style_paragraph, set_paragraph_border

# === (MỚI V26) HÀM NỘI BỘ TẠO HEADER ===
def _build_header(doc, header_data, logo_data=None):
    """Hàm nội bộ: Xây dựng Header (Tên trường, Logo, Thông tin thi)"""
    school_name = header_data.get('school_name', '').upper()
    exam_name = header_data.get('exam_name', '').upper()
    class_name = header_data.get('class_name', '').upper()
    subject_name = header_data.get('subject_name', '')
    exam_iteration = header_data.get('exam_iteration', '1')
    exam_time = header_data.get('exam_time', '90')
    
    table_header = doc.add_table(rows=1, cols=2)
    table_header.autofit = True
    
    # Cột 1: Tên trường VÀ LOGO
    cell_0 = table_header.cell(0, 0)
    cell_0.width = Cm(9) 
    p_school = cell_0.paragraphs[0]
    run_school = p_school.add_run(school_name)
    style_run(run_school, bold=True, size=12)
    style_paragraph(p_school, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1, space_after=Pt(0))
    
    if logo_data:
        try:
            p_logo = cell_0.add_paragraph()
            run_logo = p_logo.add_run()
            run_logo.add_picture(io.BytesIO(logo_data), width=Cm(2), height=Cm(2))
            style_paragraph(p_logo, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0), space_before=Pt(6))
        except Exception as e:
            print(f"Lỗi khi chèn logo vào DOCX: {e}")
    
    # Cột 2: Thông tin kỳ thi
    cell_1 = table_header.cell(0, 1)
    cell_1.width = Cm(10) 

    p_exam = cell_1.paragraphs[0]
    run_exam = p_exam.add_run(exam_name)
    style_run(run_exam, bold=True, size=12) 
    style_paragraph(p_exam, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1, space_after=Pt(0)) 
    
    p_class = cell_1.add_paragraph()
    run_class = p_class.add_run(f"LỚP: {class_name}")
    style_run(run_class, bold=True, size=12) 
    style_paragraph(p_class, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1, space_after=Pt(0)) 
    
    p_subject = cell_1.add_paragraph()
    run_subject = p_subject.add_run(f"Tên học phần: {subject_name} (Lần {exam_iteration})")
    style_run(run_subject, bold=False, size=12) 
    style_paragraph(p_subject, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1, space_after=Pt(0)) 

    p_time = cell_1.add_paragraph()
    run_time = p_time.add_run(f"Thời gian: {exam_time} phút (không kể thời gian phát đề)")
    style_run(run_time, bold=False, size=12) 
    style_paragraph(p_time, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1, space_after=Pt(0)) 

# === (MỚI V26) HÀM NỘI BỘ TẠO KHỐI KÝ TÊN ===
def _build_signoff(doc, title="Giảng viên tổng hợp đề"):
    """Hàm nội bộ: Xây dựng khối ký tên (Cần Thơ, ngày...)"""
    doc.add_paragraph() 
    
    p_signer_base = doc.add_paragraph()
    tab_stops_signer = p_signer_base.paragraph_format.tab_stops
    tab_stops_signer.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
    style_paragraph(p_signer_base, line_spacing=1.15, space_after=Pt(0))

    run_date = p_signer_base.add_run("\tCần Thơ, ngày... tháng... năm...\n")
    style_run(run_date, italic=True)
    
    run_signer = p_signer_base.add_run(f"\t{title}\n")
    style_run(run_signer, bold=True)
    
    run_name = p_signer_base.add_run("\t(Ký, ghi rõ họ tên)")
    style_run(run_name, italic=True)


# === HÀM TẠO FOOTER (LOGIC V17) ===
def create_footer(doc, total_questions):
    section = doc.sections[0]
    footer = section.footer
    
    section.footer_distance = Cm(0.5)
    section.different_first_page_header_footer = False
    
    for p in footer.paragraphs:
        p.clear()
        
    p_line = footer.add_paragraph()
    set_paragraph_border(p_line)
    style_paragraph(p_line, space_after=Pt(0), space_before=Pt(0))

    footer_table = footer.add_table(rows=1, cols=2, width=doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)
    
    cell_0 = footer_table.cell(0, 0)
    p_0 = cell_0.paragraphs[0]
    p_0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = p_0.add_run(f"Ghi chú: Đề thi gồm {total_questions} câu, được in trên ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'NUMPAGES'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    run = p_0.add_run(" trang giấy A4")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True

    cell_1 = footer_table.cell(0, 1)
    p_1 = cell_1.paragraphs[0]
    p_1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = p_1.add_run("Trang ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    run = p_1.add_run("/")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'NUMPAGES'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)


# === (SỬA V26) HÀM TẠO ĐÁP ÁN ===
def create_answer_key_doc(answer_key_map, base_name, num_tests, header_data, logo_data=None):
    doc = Document()
    
    # 1. Đặt lề (Giống file đề)
    section = doc.sections[0]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    # 2. (SỬA V26) Gọi hàm nội bộ tạo Header
    _build_header(doc, header_data, logo_data)
    
    # 3. Thêm tiêu đề "NỘI DUNG ĐÁP ÁN"
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("NỘI DUNG ĐÁP ÁN")
    style_run(run_title, bold=True, size=13)
    style_paragraph(p_title, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(10), space_before=Pt(6))

    # 4. Tạo Bảng đáp án
    num_questions = 0
    if num_tests > 0 and len(answer_key_map) > 0:
        first_test_code = list(answer_key_map.keys())[0]
        num_questions = len(answer_key_map[first_test_code])
    
    # (SỬA V26) Chia bảng theo mẫu (60 câu / 2 cột = 30 dòng)
    rows_per_col = (num_questions + 1) // 2 
    num_cols = (num_tests + 1) * 2 # (Đề\câu + 4 đề) * 2
    
    table = doc.add_table(rows=rows_per_col + 1, cols=num_cols) 
    table.style = 'Table Grid'
    table.autofit = True
    header_cells = table.rows[0].cells
    
    # Tạo Header Bảng
    for i in range(2): 
        col_offset = i * (num_tests + 1)
        header_cells[col_offset].text = 'Đề\\câu'
        for j in range(num_tests):
            header_cells[col_offset + j + 1].text = f"{base_name}{j+1:02d}"

    # Điền Bảng
    for row in range(rows_per_col):
        for col_group in range(2): 
            col_offset = col_group * (num_tests + 1)
            question_num = row + (col_group * rows_per_col) + 1
            if question_num > num_questions: break 
            table.cell(row + 1, col_offset).text = str(question_num)
            for test_idx in range(num_tests):
                test_code = f"{base_name}{test_idx+1:02d}"
                if test_code in answer_key_map and len(answer_key_map[test_code]) > (question_num - 1):
                    answer = answer_key_map[test_code][question_num - 1]
                    table.cell(row + 1, col_offset + test_idx + 1).text = answer
    
    # 5. Thêm "HẾT"
    p_het = doc.add_paragraph()
    run_het = p_het.add_run("HẾT")
    style_run(run_het, bold=True, size=13)
    style_paragraph(p_het, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6), space_before=Pt(6))
    
    # 6. (SỬA V26) Gọi hàm nội bộ tạo khối Ký tên
    _build_signoff(doc, title="Giảng viên tổng hợp đề-đáp án")
    
    # 7. (SỬA V26) Gọi hàm tạo Footer
    create_footer(doc, num_questions)
    
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# === HÀM CHÍNH TẠO FILE ZIP (SỬA V26) ===
def build_mixed_test_zip(groups, num_tests, base_name, header_data, logo_data=None):
    
    question_regex = re.compile(r"^(Câu|Question)\s+\d+[\.:]?\s+", re.IGNORECASE)
    
    answer_key_map = {}
            
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        
        for i in range(1, num_tests + 1):
            test_code = f"{base_name}{i:02d}" 
            answer_key_map[test_code] = [] 
            
            doc = Document() 
            
            section = doc.sections[0]
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            
            # (SỬA V26) 1. Gọi hàm nội bộ tạo Header
            _build_header(doc, header_data, logo_data)
            
            # 2. TẠO "ĐỀ SỐ"
            doc_text = "(HSSV không được sử dụng tài liệu)" if not header_data.get('allow_documents', False) else "(HSSV được sử dụng tài liệu)"
            p_de = doc.add_paragraph()
            run_de = p_de.add_run(f"ĐỀ SỐ: {test_code} ")
            style_run(run_de, bold=True, size=13)
            run_doc = p_de.add_run(doc_text)
            style_run(run_doc, bold=False, size=13)
            style_paragraph(p_de, line_spacing=1.15, space_after=Pt(0), keep_with_next=False, space_before=Pt(6))

            # 3. TẠO "NỘI DUNG ĐỀ THI"
            p_title = doc.add_paragraph()
            run_title = p_title.add_run("NỘI DUNG ĐỀ THI")
            style_run(run_title, bold=True, size=13)
            style_paragraph(p_title, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15, space_after=Pt(10), space_before=Pt(0), keep_with_next=False, page_break_before=False)
            
            question_counter = 1
            sorted_group_tags = sorted(groups.keys())
            
            # 4. TẠO CÂU HỎI (Loop)
            for tag in sorted_group_tags:
                question_list = groups[tag]
                
                if tag in ['g1', 'g3']:
                    random.shuffle(question_list)
                
                for q in question_list:
                    original_text = q['question_text']
                    match = question_regex.match(original_text)
                    clean_question_text = original_text.replace(match.group(0), "").strip() if match else original_text.strip()
                    
                    p_q = doc.add_paragraph()
                    style_paragraph(p_q, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, space_after=Pt(0), page_break_before=False, keep_with_next=True)
                    
                    run_prefix = p_q.add_run(f"Câu {question_counter}: ")
                    style_run(run_prefix, bold=True) 
                    
                    run_text = p_q.add_run(clean_question_text)
                    style_run(run_text, bold=False)
                    question_counter += 1
                    
                    answers = json.loads(q['answers'])
                    correct_answer_original_prefix = q['correct_answer'] 
                    
                    if tag in ['g2', 'g3']:
                        random.shuffle(answers)
                    
                    answer_prefixes = ['A', 'B', 'C', 'D']
                    found_correct_answer = False 
                    
                    table_ans = doc.add_table(rows=2, cols=2)
                    table_ans.autofit = True
                    table_ans.alignment = WD_TABLE_ALIGNMENT.CENTER 
                    
                    ans_cells = [table_ans.cell(0,0), table_ans.cell(0,1), table_ans.cell(1,0), table_ans.cell(1,1)]
                    
                    for j, ans in enumerate(answers[:4]):
                        new_prefix = answer_prefixes[j] 
                        
                        p_ans = ans_cells[j].paragraphs[0]
                        ans_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.TOP 
                        
                        style_paragraph(p_ans, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, space_after=Pt(0), page_break_before=False)
                        p_ans.paragraph_format.left_indent = Cm(0.5)
                        
                        run_p_prefix = p_ans.add_run(f"{new_prefix}. ")
                        style_run(run_p_prefix, bold=True) 
                        
                        run_p_text = p_ans.add_run(ans['text'])
                        style_run(run_text)
                        
                        if ans['prefix'] == correct_answer_original_prefix:
                            answer_key_map[test_code].append(new_prefix)
                            found_correct_answer = True

                    if not found_correct_answer:
                        answer_key_map[test_code].append('?') 

            # 5. (SỬA V26) Gọi hàm nội bộ tạo khối Ký tên
            _build_signoff(doc, title="Giảng viên tổng hợp đề")
            
            # 6. (SỬA V26) Gọi hàm tạo Footer
            create_footer(doc, question_counter - 1)

            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            file_name = f"Ma_de_{test_code}.docx"
            zip_file.writestr(file_name, doc_buffer.read())

        # 7. (SỬA V26) Gọi hàm tạo tệp đáp án (đã sửa)
        answer_key_buffer = create_answer_key_doc(answer_key_map, base_name, num_tests, header_data, logo_data)
        zip_file.writestr(f"Dap_an_Tong_hop_{base_name}.docx", answer_key_buffer.read())

    zip_buffer.seek(0)
    
    return zip_buffer
