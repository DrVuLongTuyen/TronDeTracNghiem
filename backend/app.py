import os
import re
import docx
import json
import random 
import io       
import zipfile  
from docx import Document 
from docx.shared import Pt, Cm 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime 
import traceback 
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from supabase import create_client, Client 

# --- Cấu hình ---
app = Flask(__name__)
CORS(app, expose_headers=['content-disposition', 'Content-Disposition']) 

# --- Kết nối Supabase ---
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
print("Đã kết nối Supabase (Backend Service Role)")


# --- HÀM /UPLOAD (ĐÃ SỬA LỖI ĐỌC CÂU HỎI NHIỀU DÒNG) ---

def parse_test_document(file_stream):
    try:
        doc = docx.Document(file_stream)
    except Exception as e:
        return None, f"Lỗi đọc tệp DOCX: {e}"

    test_structure = []
    current_group = None
    current_question = None
    pending_question_text = ""

    group_regex = re.compile(r"<(/?#?g\d+)>")
    question_regex = re.compile(r"^(Câu|Question)\s+\d+[\.:]?\s+", re.IGNORECASE)
    answer_regex = re.compile(r"^(#?[A-Z])[\.:]?\s+", re.IGNORECASE) 

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        group_match = group_regex.search(text)
        if group_match:
            group_tag = group_match.group(1).replace('#', '').replace('/', '')
            current_group = {"group_tag": group_tag, "is_fixed": '#' in group_tag, "questions": []}
            test_structure.append(current_group)
            current_question = None
            pending_question_text = ""
            continue
            
        question_match = question_regex.search(text)
        if question_match:
            if current_group is None:
                current_group = { "group_tag": "g3", "is_fixed": False, "questions": [] }
                test_structure.append(current_group)
            
            pending_question_text = text 
            current_question = None 
            continue
            
        answer_match = answer_regex.search(text)
        if answer_match:
            if pending_question_text and current_question is None:
                current_question = {"question_text": pending_question_text, "answers": [], "correct_answer": None}
                current_group["questions"].append(current_question)
                pending_question_text = "" 
            
            if current_question is not None:
                answer_prefix = answer_match.group(1).upper().replace('#', '')
                answer_text = text[answer_match.end():].strip()
                is_fixed = '#' in answer_match.group(0)
                is_correct = any(run.font.underline for run in para.runs)
                
                answer_obj = {"prefix": answer_prefix, "text": answer_text, "is_fixed": is_fixed}
                current_question["answers"].append(answer_obj)
                if is_correct:
                    current_question["correct_answer"] = answer_prefix
                continue
        
        if pending_question_text and not group_match and not question_match and not answer_match:
            pending_question_text += "\n" + text 
            continue

    return test_structure, None

@app.route('/')
def home():
    return "Xin chào, API Backend của TronDeTN đang hoạt động!"

@app.route('/upload', methods=['POST'])
def handle_upload():
    # (Hàm này giữ nguyên, không thay đổi)
    auth_header = request.headers.get('Authorization')
    if not auth_header: return jsonify({"error": "Thiếu token xác thực"}), 401
    try:
        jwt_token = auth_header.split(' ')[1]
        user_id = supabase.auth.get_user(jwt_token).user.id
    except Exception as e:
        return jsonify({"error": f"Token không hợp lệ: {e}"}), 401
    if 'file' not in request.files: return jsonify({"error": "Không có tệp"}), 400
    file = request.files['file']
    if not file.filename.endswith('.docx'): return jsonify({"error": "Chỉ chấp nhận .docx"}), 400
    parsed_data, error = parse_test_document(file.stream)
    if error: return jsonify({"error": error}), 500
    rows_to_insert = []
    for group in parsed_data:
        for question in group["questions"]:
            if not question["answers"]: continue
            rows_to_insert.append({
                "user_id": user_id, "group_tag": group["group_tag"], "group_is_fixed": group["is_fixed"],
                "question_text": question["question_text"], "answers": json.dumps(question["answers"], ensure_ascii=False),
                "correct_answer": question.get("correct_answer")
            })
    if not rows_to_insert: return jsonify({"error": "Không tìm thấy câu hỏi hợp lệ"}), 400
    try:
        supabase.table('question_banks').insert(rows_to_insert).execute()
    except Exception as e:
        return jsonify({"error": f"Lỗi khi lưu vào DB: {e}"}), 500
    return jsonify({"message": f"Xử lý thành công tệp '{file.filename}'", "questions_saved": len(rows_to_insert)}), 200

# --- ENDPOINT XÓA KHO CÂU HỎI (ĐÃ SỬA LỖI) ---

@app.route('/clear', methods=['DELETE'])
def handle_clear():
    # (Hàm này giữ nguyên, không thay đổi)
    auth_header = request.headers.get('Authorization')
    if not auth_header: return jsonify({"error": "Thiếu token xác thực"}), 401
    try:
        jwt_token = auth_header.split(' ')[1]
        user_id = supabase.auth.get_user(jwt_token).user.id
    except Exception as e:
        return jsonify({"error": f"Token không hợp lệ: {e}"}), 401
    try:
        response = supabase.table('question_banks').delete().eq('user_id', user_id).execute()
        num_deleted = len(response.data) 
        return jsonify({"message": f"Đã xóa thành công {num_deleted} câu hỏi cũ."}), 200
    except Exception as e:
        return jsonify({"error": f"Lỗi khi xóa dữ liệu: {e}"}), 500

    
# --- (NÂNG CẤP) HÀM TẠO ĐÁP ÁN ---

def create_answer_key_doc(answer_key_map, base_name, num_tests):
    # (Hàm này giữ nguyên, không thay đổi)
    doc = Document()
    doc.add_heading("NỘI DUNG ĐÁP ÁN", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() 
    num_questions = 0
    if num_tests > 0 and len(answer_key_map) > 0:
        first_test_code = list(answer_key_map.keys())[0]
        num_questions = len(answer_key_map[first_test_code])
    rows_per_col = (num_questions + 1) // 2 
    num_cols = (num_tests + 1) * 2
    table = doc.add_table(rows=rows_per_col + 1, cols=num_cols) 
    table.style = 'Table Grid'
    table.autofit = True
    header_cells = table.rows[0].cells
    for i in range(2): 
        col_offset = i * (num_tests + 1)
        header_cells[col_offset].text = 'Đề\\câu'
        for j in range(num_tests):
            header_cells[col_offset + j + 1].text = f"{base_name}{j+1:02d}"
    for row in range(rows_per_col):
        for col_group in range(2): 
            col_offset = col_group * (num_tests + 1)
            question_num = row + (col_group * rows_per_col) + 1
            if question_num > num_questions: break 
            table.cell(row + 1, col_offset).text = str(question_num)
            for test_idx in range(num_tests):
                test_code = f"{base_name}{test_idx+1:02d}"
                if test_code in answer_key_map and len(answer_key_map[test_code]) > (question_num - 1):
                    answer = answer_key_map[test_code][question_num - 1]
                    table.cell(row + 1, col_offset + test_idx + 1).text = answer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# --- (MỚI) HÀM TẠO ĐƯỜNG KẺ NGANG (BORDER) CHO PARAGRAPH ---
def set_paragraph_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr() 
    pBdr = OxmlElement('w:pBdr')       
    
    topBdr = OxmlElement('w:top')
    topBdr.set(qn('w:val'), 'single') 
    topBdr.set(qn('w:sz'), '4')       
    topBdr.set(qn('w:space'), '1')
    topBdr.set(qn('w:color'), 'auto')
    
    pBdr.append(topBdr)
    pPr.append(pBdr)

# --- (SỬA LỖI) HÀM TẠO FOOTER (CHÂN TRANG) ---
def create_footer(doc, total_questions):
    section = doc.sections[0]
    footer = section.footer
    
    # (SỬA) Đẩy footer xuống thấp
    section.footer_distance = Cm(0.5)
    
    # Xóa mọi thứ cũ trong footer
    for p in footer.paragraphs:
        p.clear()
        
    # (SỬA) 1. Thêm đường kẻ ngang LÊN TRÊN
    p_line = footer.add_paragraph()
    set_paragraph_border(p_line)
    # (SỬA LỖI) Thêm dãn đoạn 0pt cho đường kẻ
    style_paragraph(p_line, space_after=Pt(2), space_before=0) # Thêm 2pt space after
    
    # 2. Thêm bảng 2 cột BÊN DƯỚI đường kẻ
    footer_table = footer.add_table(rows=1, cols=2, width=doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)
    
    # Cột 1: Ghi chú
    cell_0 = footer_table.cell(0, 0)
    p_0 = cell_0.paragraphs[0]
    p_0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p_0, line_spacing=1, space_after=0, space_before=0)
    
    # (MỚI) Thêm trường NUMPAGES (Tổng số trang) vào Ghi chú
    run = p_0.add_run(f"Ghi chú: Đề thi gồm {total_questions} câu, được in trên ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    
    # Thêm field code cho TỔNG SỐ TRANG
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'NUMPAGES'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    run = p_0.add_run(" trang giấy A4")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True


    # Cột 2: Số trang
    cell_1 = footer_table.cell(0, 1)
    p_1 = cell_1.paragraphs[0]
    p_1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p_1, line_spacing=1, space_after=0, space_before=0)
    
    # Thêm field code cho TRANG HIỆN TẠI
    run = p_1.add_run("Trang ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    run = p_1.add_run("/")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

    # Thêm field code cho TỔNG SỐ TRANG
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'NUMPAGES'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
# --- (MỚI) HÀM STYLE CHUNG ---
def style_run(run, bold=False, italic=False, size=13):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    
# (SỬA LỖI) Sửa lại hàm style_paragraph để gọi đúng (space_before ở cuối)
def style_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15, space_after=0, page_break_before=False, keep_with_next=False, space_before=0):
    p.paragraph_format.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before) 
    p.paragraph_format.page_break_before = page_break_before 
    p.paragraph_format.keep_with_next = keep_with_next 
    p.paragraph_format.widow_control = False # Tắt widow control để fix lỗi ngắt trang

# --- (NÂNG CẤP LỚN) ENDPOINT TRỘN ĐỀ (GIAI ĐOẠN 3) ---

@app.route('/mix', methods=['POST'])
def handle_mix():
    # (SỬA LỖI) Thêm dòng định nghĩa question_regex bị thiếu
    question_regex = re.compile(r"^(Câu|Question)\s+\d+[\.:]?\s+", re.IGNORECASE)
        
    try:
        # 1. Xác thực người dùng
        auth_header = request.headers.get('Authorization')
        if not auth_header: return jsonify({"error": "Thiếu token xác thực"}), 401
        
        try:
            jwt_token = auth_header.split(' ')[1]
            user_id = supabase.auth.get_user(jwt_token).user.id
        except Exception as e:
            return jsonify({"error": f"Token không hợp lệ: {e}"}), 401

        # 2. Lấy thông số từ Frontend
        data = request.json
        num_tests = int(data.get('num_tests', 2))
        base_name = data.get('base_name', 'VLT').upper()
        header_data = data.get('header_data', {})
        school_name = header_data.get('school_name', '').upper()
        exam_name = header_data.get('exam_name', '').upper()
        class_name = header_data.get('class_name', '').upper()
        subject_name = header_data.get('subject_name', '')
        exam_iteration = header_data.get('exam_iteration', '1')
        exam_time = header_data.get('exam_time', '90')
        allow_documents = header_data.get('allow_documents', False)
        
        # 3. Lấy toàn bộ câu hỏi của user từ DB
        try:
            response = supabase.table('question_banks').select("*").eq('user_id', user_id).execute()
            if not response.data:
                return jsonify({"error": "Không tìm thấy câu hỏi nào trong kho dữ liệu. Vui lòng upload đề trước."}), 404
            
            groups = {}
            for q in response.data:
                tag = q['group_tag']
                if tag not in groups:
                    groups[tag] = []
                groups[tag].append(q)
                
        except Exception as e:
            return jsonify({"error": f"Lỗi khi lấy dữ liệu từ DB: {e}"}), 500
            
        answer_key_map = {}
            
        # 4. Tạo file Zip trong bộ nhớ
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # 5. Lặp để tạo số lượng đề
            for i in range(1, num_tests + 1):
                test_code = f"{base_name}{i:02d}" 
                answer_key_map[test_code] = [] 
                
                doc = Document() 
                
                # --- (SỬA LỖI) SET LỀ (MARGINS) 1CM ---
                section = doc.sections[0]
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                
                # --- (SỬA LỖI) TẠO HEADER THEO MẪU (VỚI TAB 15CM) ---
                table_header = doc.add_table(rows=1, cols=2)
                table_header.autofit = True
                
                # Cột 1: Tên trường (Căn giữa)
                cell_0 = table_header.cell(0, 0)
                cell_0.width = Cm(9) 
                p_school = cell_0.paragraphs[0]
                run_school = p_school.add_run(school_name)
                style_run(run_school, bold=True, size=12) # (SỬA) Size 12
                # (SỬA LỖI) Gọi hàm style_paragraph với keyword arguments
                style_paragraph(p_school, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1, space_after=0)
                
                # Cột 2: Thông tin kỳ thi (Tab 15cm)
                cell_1 = table_header.cell(0, 1)
                cell_1.width = Cm(10) 

                # (SỬA LỖI) Sửa lại logic Căn lề phải (theo mẫu image_92e8c6.png)
                p_exam = cell_1.paragraphs[0]
                run_exam = p_exam.add_run(exam_name)
                style_run(run_exam, bold=True, size=12) 
                style_paragraph(p_exam, align=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1, space_after=0) # (SỬA) Căn trái
                
                p_class = cell_1.add_paragraph()
                run_class = p_class.add_run(f"LỚP: {class_name}")
                style_run(run_class, bold=True, size=12) 
                style_paragraph(p_class, align=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1, space_after=0) # (SỬA) Căn trái
                
                p_subject = cell_1.add_paragraph()
                run_subject = p_subject.add_run(f"Tên học phần: {subject_name} (Lần {exam_iteration})")
                style_run(run_subject, bold=False, size=12) 
                style_paragraph(p_subject, align=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1, space_after=0) # (SỬA) Căn trái

                p_time = cell_1.add_paragraph()
                run_time = p_time.add_run(f"Thời gian: {exam_time} phút (không kể thời gian phát đề)")
                style_run(run_time, bold=False, size=12) 
                style_paragraph(p_time, align=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1, space_after=0) # (SỬA) Căn trái

                doc.add_paragraph() 

                # --- (MỚI) TẠO THÔNG TIN ĐỀ SỐ ---
                doc_text = "(HSSV không được sử dụng tài liệu)" if not allow_documents else "(HSSV được sử dụng tài liệu)"
                p_de = doc.add_paragraph()
                run_de = p_de.add_run(f"ĐỀ SỐ: {test_code} ")
                style_run(run_de, bold=True, size=13)
                run_doc = p_de.add_run(doc_text)
                style_run(run_doc, bold=False, size=13)
                style_paragraph(p_de, line_spacing=1.15, space_after=0)

                doc.add_paragraph() 

                # --- (MỚI) TẠO TIÊU ĐỀ "NỘI DUNG" ---
                p_title = doc.add_paragraph()
                run_title = p_title.add_run("NỘI DUNG ĐỀ THI")
                style_run(run_title, bold=True, size=13)
                # (SỬA LỖI) Tắt ngắt trang (dùng keyword arguments)
                style_paragraph(p_title, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15, space_after=Pt(10), page_break_before=False, keep_with_next=True)
                
                question_counter = 1
                sorted_group_tags = sorted(groups.keys())
                
                for tag in sorted_group_tags:
                    question_list = groups[tag]
                    
                    if tag in ['g1', 'g3']:
                        random.shuffle(question_list)
                    
                    for q in question_list:
                        # (SỬA LỖI) Cắt bỏ "Câu X:" một cách an toàn
                        original_text = q['question_text']
                        match = question_regex.match(original_text)
                        clean_question_text = original_text.replace(match.group(0), "").strip() if match else original_text.strip()
                        
                        p_q = doc.add_paragraph()
                        # (SỬA LỖI) Căn đều + Tắt ngắt trang (dùng keyword arguments)
                        style_paragraph(p_q, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, space_after=0, page_break_before=False, keep_with_next=True)
                        
                        run_prefix = p_q.add_run(f"Câu {question_counter}: ")
                        style_run(run_prefix, bold=True) 
                        
                        run_text = p_q.add_run(clean_question_text)
                        style_run(run_text, bold=False)
                        question_counter += 1
                        
                        answers = json.loads(q['answers'])
                        correct_answer_original_prefix = q['correct_answer'] 
                        
                        if tag in ['g2', 'g3']:
                            random.shuffle(answers)
                        
                        answer_prefixes = ['A', 'B', 'C', 'D']
                        found_correct_answer = False 
                        
                        table_ans = doc.add_table(rows=2, cols=2)
                        table_ans.autofit = True
                        table_ans.alignment = WD_TABLE_ALIGNMENT.CENTER 
                        
                        ans_cells = [table_ans.cell(0,0), table_ans.cell(0,1), table_ans.cell(1,0), table_ans.cell(1,1)]
                        
                        for j, ans in enumerate(answers[:4]):
                            new_prefix = answer_prefixes[j] 
                            
                            p_ans = ans_cells[j].paragraphs[0]
                            ans_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.TOP 
                            
                            # (SỬA LỖI) Căn đều (dùng keyword arguments)
                            style_paragraph(p_ans, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, space_after=0, page_break_before=False)
                            p_ans.paragraph_format.left_indent = Cm(0.5)
                            
                            run_p_prefix = p_ans.add_run(f"{new_prefix}. ")
                            style_run(run_p_prefix, bold=True) 
                            
                            run_p_text = p_ans.add_run(ans['text'])
                            style_run(run_p_text)
                            
                            if ans['prefix'] == correct_answer_original_prefix:
                                answer_key_map[test_code].append(new_prefix)
                                found_correct_answer = True

                        if not found_correct_answer:
                            answer_key_map[test_code].append('?') 

                # --- (SỬA LỖI) TẠO KHỐI KÝ TÊN (Tab 14cm) ---
                doc.add_paragraph() 
                
                # Căn giữa tại 14cm
                p_signer_base = doc.add_paragraph()
                tab_stops_signer = p_signer_base.paragraph_format.tab_stops
                tab_stop_signer = tab_stops_signer.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
                style_paragraph(p_signer_base, line_spacing=1.15, space_after=0)

                run_date = p_signer_base.add_run("\tCần Thơ, ngày... tháng... năm...\n")
                style_run(run_date, italic=True)
                
                run_signer = p_signer_base.add_run("\tGiảng viên tổng hợp đề\n")
                style_run(run_signer, bold=True)
                
                run_name = p_signer_base.add_run("\t(Ký, ghi rõ họ tên)")
                style_run(run_name, italic=True)
                
                
                # --- (MỚI) TẠO FOOTER ---
                create_footer(doc, question_counter - 1)

                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                file_name = f"Ma_de_{test_code}.docx"
                zip_file.writestr(file_name, doc_buffer.read())

            # 6. Tạo tệp đáp án tổng hợp
            answer_key_buffer = create_answer_key_doc(answer_key_map, base_name, num_tests)
            zip_file.writestr(f"Dap_an_Tong_hop_{base_name}.docx", answer_key_buffer.read())

        zip_buffer.seek(0)
        
        current_time = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'Bo_de_tron_{base_name}_{current_time}.zip' 
        )
        
    except Exception as e:
        print("--- LỖI NGHIÊM TRỌNG TRONG /MIX ---")
        print(traceback.format_exc()) 
        print("---------------------------------")
        return jsonify({"error": f"Lỗi máy chủ nội bộ: {str(e)}. Vui lòng kiểm tra lại file .docx (ví dụ: thiếu đáp án, sai định dạng) hoặc liên hệ hỗ trợ."}), 500


# Chạy ứng dụng (cho Render)
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)))
